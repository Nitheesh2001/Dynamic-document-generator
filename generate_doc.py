import json
from docx import Document
from docx.shared import Pt  # Import for font size handling

def format_run(run, bold=False, underline=False, size=None):
    """
    Formats a given run with bold, underline, and font size.
    """
    if bold:
        run.bold = True
    if underline:
        run.underline = True
    if size:
        run.font.size = Pt(size)

def replace_placeholder(paragraph, placeholder, value):
    """
    Replaces a placeholder in a paragraph with the corresponding value,
    ensuring that text spanning multiple runs is handled properly.
    """
    full_text = paragraph.text
    
    if placeholder.lower() == "case_type":
        case_placeholder_variations = ["[[Case Type]]", "[[case type]]", "[[CASE TYPE]]"]
        for variation in case_placeholder_variations:
            full_text = full_text.replace(variation, str(value))
    else:
        variations = [
            f"[[{placeholder}]]",
            f"[[{placeholder.lower()}]]",
            f"[[{placeholder.upper()}]]",
            f"[[{placeholder.capitalize()}]]",
        ]
        for variation in variations:
            full_text = full_text.replace(variation, str(value))

    if paragraph.text != full_text:
        for run in paragraph.runs:
            run.text = ""
        paragraph.add_run(full_text)

def insert_exhibits(doc, exhibits):
    """
    Inserts a formatted exhibits list into the document and ensures [[exhibits]] is removed.
    """
    for paragraph in doc.paragraphs:
        if "[[exhibits]]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[[exhibits]]", "").strip()

    doc.add_paragraph('----------------------------', style='Normal')
    for index, exhibit in enumerate(exhibits, start=1):
        exhibit_name = f"Exhibit {chr(64 + index)}"  # Exhibit A, Exhibit B, etc.
        exhibit_desc = list(exhibit.values())[0]

        p = doc.add_paragraph()
        run = p.add_run(exhibit_name)
        format_run(run, bold=True, underline=True, size=12)  # Ensure exhibits are underlined and bold
        p.add_run(f"  {exhibit_desc}")  # Keep description unbolded

def replace_placeholders_in_doc(doc, data):
    """
    Replaces all placeholders in the document with corresponding values from data.
    """
    for paragraph in doc.paragraphs:
        if "Form I-129, Petitioner for Nonimmigrant Worker (H-1B Extension)" in paragraph.text:
            paragraph.clear()  # Clear existing content
            run = paragraph.add_run("Form I-129, Petitioner for Nonimmigrant Worker (H-1B Extension)")
            format_run(run, bold=True, underline=True, size=20)

        elif "Petitioner:" in paragraph.text:
            for run in paragraph.runs:
                format_run(run, bold=True, size=16)

        elif "Beneficiary:" in paragraph.text:
            for run in paragraph.runs:
                format_run(run, bold=True, size=16)
                if "Beneficiary:" not in run.text:
                    run.underline = True  # Underline only the name after "Beneficiary:"

        for placeholder, value in data.items():
            if placeholder == "exhibits":
                continue
            replace_placeholder(paragraph, placeholder, value)

def generate_document(input_path, output_path, data_path):
    """
    Reads the input Word document, replaces placeholders, and inserts exhibits.
    """
    doc = Document(input_path)

    with open(data_path, 'r') as f:
        data = json.load(f)

    replace_placeholders_in_doc(doc, data)

    if "exhibits" in data and data["exhibits"]:
        insert_exhibits(doc, data["exhibits"])

    doc.save(output_path)

if __name__ == "__main__":
    input_doc = "input.docx"
    output_doc = "output.docx"
    data_json = "data.json"

    generate_document(input_doc, output_doc, data_json)
    print(f"Document generated successfully as {output_doc}")
